"""Sinh tài liệu Word mô tả kiến trúc vận hành Sell-In Vikoda.

Tài liệu được tạo từ kiến trúc chính thức:
SharePoint -> Microsoft Graph -> GitHub Actions -> Python ETL -> SharePoint/Data_Goc.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)


def _add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "CÔNG TY CỔ PHẦN NƯỚC KHOÁNG KHÁNH HÒA — VIKODA 1979\n"
        "QUY TRÌNH TỰ ĐỘNG HÓA BÁO CÁO SELL-IN"
    )
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Phiên bản Cloud 2026 | Cập nhật: {date.today().isoformat()}")


def _add_architecture(doc: Document) -> None:
    doc.add_heading("1. Kiến trúc hệ thống", level=1)
    doc.add_paragraph(
        "Hệ thống sử dụng SharePoint Online làm nguồn dữ liệu nghiệp vụ, Microsoft Graph "
        "làm lớp đồng bộ, GitHub Actions làm bộ lập lịch/máy chạy và Python làm ETL. "
        "Các workbook đã xử lý được ghi trở lại SharePoint/Data_Goc."
    )

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Thành phần", "Vai trò", "Ghi chú"]
    for index, text in enumerate(headers):
        table.rows[0].cells[index].text = text

    rows = [
        ("SharePoint Online", "Nguồn và đích dữ liệu", "Data ERP, Target, Danh mục, Data_Goc"),
        ("Microsoft Entra ID", "Xác thực ứng dụng", "OAuth2 Client Credentials"),
        ("Microsoft Graph", "Đọc/ghi SharePoint", "Application permission Sites.Selected"),
        ("GitHub Actions", "Lịch chạy và CI", "18:00 VN hoặc manual dispatch"),
        ("Python ETL", "Chuẩn hóa và tạo báo cáo", "Strict mode + validation"),
        ("Dashboard", "Đầu ra điều hành", "Chỉ publish khi dữ liệu đã được phê duyệt"),
    ]
    for component, role, note in rows:
        cells = table.add_row().cells
        cells[0].text = component
        cells[1].text = role
        cells[2].text = note


def _add_process(doc: Document) -> None:
    doc.add_heading("2. Quy trình vận hành", level=1)
    steps = [
        "Sales Admin/Kế toán upload workbook ERP vào SharePoint/Data ERP.",
        "GitHub Actions chạy theo lịch hoặc được quản trị viên chạy thủ công.",
        "Workflow lấy token từ Microsoft Entra ID và resolve SharePoint site/drive.",
        "Microsoft Graph tải Data ERP, Target, Danh mục khách hàng và Danh mục sản phẩm.",
        "run_cloud_pipeline.py --strict chạy ETL, tạo workbook tháng, báo cáo tổng hợp và web data.",
        "health_check.py kiểm tra chất lượng; thiếu dữ liệu hoặc artifact mới thì workflow dừng.",
        "Các workbook Data/Data_Goc được upload trở lại SharePoint/Data_Goc qua Microsoft Graph.",
        "Dashboard chỉ được publish khi ENABLE_PAGES_DEPLOY=true và dữ liệu được đánh dấu public-or-sanitized.",
    ]
    for index, step in enumerate(steps, start=1):
        doc.add_paragraph(f"Bước {index}: {step}")


def _add_configuration(doc: Document) -> None:
    doc.add_heading("3. Cấu hình bắt buộc", level=1)
    doc.add_paragraph("GitHub Repository Secrets:")
    for name in ("AZURE_TENANT_ID", "AZURE_CLIENT_ID", "AZURE_CLIENT_SECRET"):
        doc.add_paragraph(name, style="List Bullet")

    doc.add_paragraph("Repository Variables:")
    for name in (
        "SHAREPOINT_HOSTNAME",
        "SHAREPOINT_SITE_PATH",
        "SHAREPOINT_ERP_FOLDER",
        "SHAREPOINT_TARGET_FOLDER",
        "SHAREPOINT_CUSTOMER_FOLDER",
        "SHAREPOINT_PRODUCT_FOLDER",
        "SHAREPOINT_DATA_GOC_FOLDER",
    ):
        doc.add_paragraph(name, style="List Bullet")


def _add_security(doc: Document) -> None:
    doc.add_heading("4. Kiểm soát bảo mật", level=1)
    controls = [
        "Không commit dữ liệu ERP, workbook sản xuất, .env hoặc secret lên Git.",
        "App Entra chỉ được cấp quyền cần thiết trên site Planning.",
        "CI pull request/push không dùng production secrets.",
        "Cloud pipeline fail-closed khi thiếu credential hoặc thiếu workbook nguồn.",
        "Static hosting không được dùng cho dữ liệu nội bộ chưa sanitize.",
    ]
    for control in controls:
        doc.add_paragraph(control, style="List Bullet")


def create_document(output_path: Path) -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    _set_default_font(doc)
    _add_title(doc)
    _add_architecture(doc)
    _add_process(doc)
    _add_configuration(doc)
    _add_security(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    project_root = Path(__file__).resolve().parents[4]
    target = project_root / "Quy_Trinh_Tu_Dong_Hoa_Bao_Cao_SellIn_Vikoda.docx"
    created = create_document(target)
    print(f"Đã tạo tài liệu: {created}")
